from __future__ import annotations
from pathlib import Path
import zipfile
import docx  # python-docx

from core.registry import register
from core.schemas import ParsedDoc, TextBlock, ImageAsset, TableAsset
from core.parsing.base import BaseParser


@register("parsing", "docx")
class DocxParser(BaseParser):
    def parse(self, path: Path, doc_id: str) -> ParsedDoc:
        d = docx.Document(path)

        text_blocks: list[TextBlock] = []
        section_path = ""
        for para in d.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading"):
                section_path = para.text.strip()
                continue
            text_blocks.append(TextBlock(text=para.text.strip(), section_path=section_path))

        tables: list[TableAsset] = []
        for i, table in enumerate(d.tables, start=1):
            rows = ["| " + " | ".join(c.text.strip() for c in row.cells) + " |" for row in table.rows]
            tables.append(TableAsset(markdown="\n".join(rows), id=f"{doc_id}_tbl{i}"))

        images: list[ImageAsset] = []
        if self.extract_images:
            # A .docx is a zip; embedded images live under word/media/.
            with zipfile.ZipFile(path) as z:
                media = [n for n in z.namelist() if n.startswith("word/media/")]
                for i, name in enumerate(media, start=1):
                    img_id = f"{doc_id}_img{i}"
                    ext = name.rsplit(".", 1)[-1]
                    out_path = self.assets_dir / f"{img_id}.{ext}"
                    out_path.write_bytes(z.read(name))
                    images.append(ImageAsset(path=out_path, id=img_id))

        return ParsedDoc(
            doc_id=doc_id, source_path=path, format="docx",
            text_blocks=text_blocks, images=images, tables=tables, n_pages=1,
            meta={"parser": "python-docx"},
        )
